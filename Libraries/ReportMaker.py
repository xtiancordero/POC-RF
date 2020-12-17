from docx import Document
from docx.shared import Inches
from PIL import ImageGrab
from ctypes import windll
from PIL import Image

# document = Document()
# document = Document(f)
# f = open(fileName + '.docx', 'rb')
# toplist, winlist = [], []

class ReportMaker(object):
    def __init__(self):
        self.mf = Document()
        # self.toplist = []
        # self.winlist = []

    def set_report_header(self, reportHeader):
        # self.mf.add_heading(reportHeader, 0)
        # self.mf._body.clear_content()
        self.mf.add_heading(reportHeader, 0)

    def set_report_sub_header(self, reportHeader):
        self.mf.add_heading(reportHeader, level=1)

    def put_screenshot(self, path, name):
        pathHolder = path + '.png'
        self.mf.add_picture(pathHolder, width=Inches(6), height=Inches(2.9))
        self.mf.add_page_break()

    def put_text(self, text1):
        self.mf.add_paragraph().add_run(text1)

    def put_screenshot_with_text(self, path, name):
        self.mf.add_paragraph().add_run(name)
        pathHolder = path + '.png'
        self.mf.add_picture(pathHolder, width=Inches(6), height=Inches(2.9))
        self.mf.add_page_break()

    def add_log_in_document(self, log, isHeader=False):
        if str(isHeader).lower() == 'true':
            self.mf.add_heading(log, level=2)
        else:
            self.mf.add_paragraph().add_run(log)

    def fail_screenshot(self, error_Message, path):
        pathHolder = path + '.png'
        self.mf.add_heading("FAILED: " + error_Message, level=2)
        self.mf.add_picture(pathHolder, width=Inches(6), height=Inches(2.9))

    def save_document(self, fileName):
        self.mf.save(fileName + '.docx')

    # def get_image(self, path):
    #     name = ""
    #     hwnd = win32gui.FindWindow(None, name)
    #     left, top, right, bot = win32gui.GetWindowRect(hwnd)
    #     w = right - left
    #     h = bot - top
	#
    #     hwndDC = win32gui.GetWindowDC(hwnd)
    #     mfcDC  = win32ui.CreateDCFromHandle(hwndDC)
    #     saveDC = mfcDC.CreateCompatibleDC()
	#
    #     saveBitMap = win32ui.CreateBitmap()
    #     saveBitMap.CreateCompatibleBitmap(mfcDC, w, h)
	#
    #     saveDC.SelectObject(saveBitMap)
    #     result = windll.user32.PrintWindow(hwnd, saveDC.GetSafeHdc(), 0)
    #     print (result)
	#
    #     bmpinfo = saveBitMap.GetInfo()
    #     bmpstr = saveBitMap.GetBitmapBits(True)
	#
    #     im = Image.frombuffer(
    #         'RGB',
    #         (bmpinfo['bmWidth'], bmpinfo['bmHeight']),
    #         bmpstr, 'raw', 'BGRX', 0, 1)
	#
    #     win32gui.DeleteObject(saveBitMap.GetHandle())
    #     saveDC.DeleteDC()
    #     mfcDC.DeleteDC()
    #     win32gui.ReleaseDC(hwnd, hwndDC)
	#
    #     if result == 1:
    #         im.save(path)
	#
